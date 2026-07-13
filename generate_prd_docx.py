#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成北京大学学院看板移动端PRD Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def create_doc():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    
    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Microsoft YaHei'
        heading_style.font.color.rgb = RGBColor(0x8C, 0x15, 0x15)  # 北大红
    
    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('北京大学学院看板移动端')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8C, 0x15, 0x15)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('产品需求文档（PRD）')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\n')
    
    # 文档信息表
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info = [
        ('产品名称', '北京大学学院看板移动端'),
        ('版本', 'V1.0'),
        ('文档状态', '初稿'),
        ('创建日期', '2026-07-13'),
        ('目标用户', '二级学院院长'),
    ]
    
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_page_break()
    
    # 目录页
    doc.add_heading('目录', level=1)
    toc_items = [
        '一、产品概述',
        '二、用户画像与使用场景',
        '三、功能需求',
        '四、信息架构',
        '五、交互规范',
        '六、视觉规范',
        '七、数据需求',
        '八、非功能需求',
        '九、版本规划',
        '十、附录',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # 一、产品概述
    doc.add_heading('一、产品概述', level=1)
    
    doc.add_heading('1.1 产品背景', level=2)
    doc.add_paragraph(
        '北京大学二级学院院长需要及时掌握本学院人事人才全貌，包括教职工结构、'
        '人才分布、人员流动等关键数据，以支撑日常管理和战略决策。目前缺乏便捷的'
        '移动端数据查看工具，院长需要通过PC端或人工汇报获取信息，效率较低。'
    )
    
    doc.add_heading('1.2 产品定位', level=2)
    doc.add_paragraph(
        '面向二级学院院长的移动端人事数据看板，提供实时、直观、可穿透的人事数据'
        '分析，帮助院长快速掌握学院人才全貌，支持管理决策。'
    )
    
    doc.add_heading('1.3 核心价值', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    values = [
        ('价值点', '说明'),
        ('随时随地', '手机端随时查看，不依赖PC'),
        ('一目了然', '核心KPI一眼掌握'),
        ('数据穿透', '从概览到明细逐层深入'),
        ('横向对比', '了解本学院在全校的相对位置'),
    ]
    for i, (k, v) in enumerate(values):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading('1.4 技术方案', level=2)
    tech_items = [
        '形态：移动端H5页面',
        '适配：iOS/Android手机浏览器',
        '访问方式：通过北大企业微信或链接访问',
        '数据更新：每日自动同步人事数据',
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 二、用户画像与使用场景
    doc.add_heading('二、用户画像与使用场景', level=1)
    
    doc.add_heading('2.1 用户画像', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    users = [
        ('角色', '特征', '核心需求'),
        ('二级学院院长', '45-60岁，管理学院日常事务', '快速了解学院人事全貌，支撑决策'),
        ('学院书记', '同上', '关注党建、人才引进'),
        ('分管副院长', '协助院长管理', '关注分管领域数据'),
    ]
    for i, row_data in enumerate(users):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('2.2 核心使用场景', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    scenarios = [
        ('场景', '频率', '用户目标', '设计重点'),
        ('晨间速览', '每日', '快速了解学院最新动态', '首页KPI一目了然'),
        ('上级汇报', '周/月', '向校领导汇报学院情况', '数据可截图、可导出'),
        ('人才引进决策', '不定期', '评估现有人才结构缺口', '人才分布、学缘分析'),
        ('预算编制规划', '年度', '评估人员编制需求', '拟退预警、进出趋势'),
        ('学院对比', '不定期', '了解学院在全校的位置', '匿名化排名/区间对比'),
    ]
    for i, row_data in enumerate(scenarios):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_page_break()
    
    # 三、功能需求
    doc.add_heading('三、功能需求', level=1)
    
    doc.add_heading('3.1 功能架构', level=2)
    arch_text = """学院看板移动端
├── 首页总览
│   ├── 核心KPI卡片
│   ├── 分布概览入口
│   └── 学院排名摘要
├── 分析详情
│   ├── 年龄分布
│   ├── 性别分布
│   ├── 政治面貌分布
│   ├── 学历结构
│   ├── 职称结构
│   ├── 学缘结构
│   ├── 人员进出
│   └── 拟退预警
├── 人员名单
│   ├── 在职人员名单
│   ├── 合同制人员名单
│   ├── 博士后名单
│   ├── 离退人员名单
│   └── 人才名单
└── 设置
    └── 基本设置"""
    p = doc.add_paragraph(arch_text)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    
    doc.add_heading('3.2 功能模块详细说明', level=2)
    
    # 模块1：人员概况
    doc.add_heading('模块1：人员概况', level=3)
    doc.add_paragraph('功能描述：展示学院各类人员数量及占比，支持全校对比排名。')
    
    doc.add_paragraph('展示内容：')
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    data = [
        ('指标', '说明', '数据来源'),
        ('在职总人数', '含专任教师、职员、其他', '人员基本信息宽表'),
        ('专任教师数', '教学科研人员', '人员基本信息宽表'),
        ('职员数', '行政管理、教辅人员', '人员基本信息宽表'),
        ('合同制人数', '合同制、派遣制', '合同制职工信息'),
        ('博士后人数', '在站博士后', '博士后信息'),
        ('离退人数', '已退休人员', '人员基本信息宽表'),
    ]
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph('交互说明：')
    interactions = [
        '点击任意数字可下钻到对应人员名单',
        '每个数字旁显示全校排名徽章（如"第3/28"）',
        '点击排名可查看全校各学院对比',
    ]
    for item in interactions:
        doc.add_paragraph(item, style='List Bullet')
    
    # 模块2-9 简化展示
    modules = [
        ('模块2：年龄分布', '展示学院人员年龄结构，支持按人员类别筛选。包含平均年龄、年龄段分布柱状图、全校排名、近3年趋势。'),
        ('模块3：性别分布', '展示学院人员性别结构，支持全校对比。包含男女比例环形图、人数占比、全校排名。'),
        ('模块4：政治面貌分布', '展示学院人员政治面貌结构。包含中共党员、民主党派、群众人数占比、全校排名。'),
        ('模块5：学历结构', '展示学院人员学历分布，重点关注博士占比。包含博士/硕士/本科占比、全校排名。'),
        ('模块6：职称结构', '展示学院人员职称分布，区分事业编和新体制。包含正高/副高/中级/初级占比、全校排名。'),
        ('模块7：学缘结构', '展示学院人员学缘来源，重点关注北大学缘和海外学缘。包含本校/国内/海外占比、TOP10院校。'),
        ('模块8：人员进出', '展示学院近期人员流动情况。包含近30天入职/减离人数、近12个月趋势、全校排名。'),
        ('模块9：拟退预警', '展示未来拟退休人员情况。包含未来1年/3年拟退人数、按职称/学科分布、全校排名。'),
        ('模块10：人才情况', '展示学院高层次人才信息。包含院士、长江学者、杰青、青年人才人数及名单。'),
        ('模块11：获奖情况', '展示学院人员获奖信息。包含获奖总人次、按年度/级别分布、获奖名单。'),
    ]
    
    for title, desc in modules:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)
    
    # 人员名单功能
    doc.add_heading('3.3 人员名单功能', level=3)
    doc.add_paragraph('功能描述：展示各类人员的详细名单，支持搜索和筛选。')
    
    doc.add_paragraph('展示字段：')
    fields = ['姓名、职工号、职称/职位、最高学历、人员类别、联系方式']
    for item in fields:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph('交互说明：')
    list_interactions = [
        '支持按姓名搜索',
        '支持按职称、学历、人员类别筛选',
        '点击单条可展开查看详情',
        '支持点击手机号拨打电话',
    ]
    for item in list_interactions:
        doc.add_paragraph(item, style='List Bullet')
    
    # 全校对比功能
    doc.add_heading('3.4 全校对比功能', level=3)
    doc.add_paragraph('功能描述：展示本学院各指标在全校的相对位置。')
    
    doc.add_paragraph('展示形式：')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    compare = [
        ('形式', '说明'),
        ('排名徽章', '"第3/28"显示排名'),
        ('区间分布', '"高于75%的学院"显示相对位置'),
        ('柱状图对比', '可视化展示各学院排名（匿名化）'),
    ]
    for i, (k, v) in enumerate(compare):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_paragraph('数据权限说明：')
    permissions = [
        '院长可查看本学院所有数据明细',
        '全校对比仅展示排名和区间，不展示其他学院具体数值',
        '不支持查看其他学院的人员名单',
    ]
    for item in permissions:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 四、信息架构
    doc.add_heading('四、信息架构', level=1)
    
    doc.add_heading('4.1 页面层级', level=2)
    layers = [
        'L1 首页总览：核心KPI卡片、分布概览入口、学院排名摘要',
        'L2 分析详情：各维度分布图表、全校对比排名、趋势图',
        'L3 人员名单：可搜索名单、个人详情卡片、拨打电话',
    ]
    for item in layers:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.2 导航结构', level=2)
    doc.add_paragraph('顶部导航栏：左侧北大校徽+校名+学院名称，右侧设置图标')
    doc.add_paragraph('底部Tab导航：首页 | 分析 | 名单 | 设置')
    
    doc.add_page_break()
    
    # 五、交互规范
    doc.add_heading('五、交互规范', level=1)
    
    doc.add_heading('5.1 导航交互', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    nav = [
        ('操作', '说明'),
        ('左滑返回', '支持左滑手势返回上一页'),
        ('返回按钮', '左上角返回箭头'),
        ('底部Tab', '切换主要功能模块'),
    ]
    for i, (k, v) in enumerate(nav):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading('5.2 数据交互', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data_int = [
        ('操作', '说明'),
        ('下拉刷新', '支持下拉刷新数据'),
        ('点击穿透', '图表/数字点击可下钻到明细'),
        ('筛选器', '支持按人员类别筛选'),
        ('搜索', '支持按姓名/职称/学历搜索'),
    ]
    for i, (k, v) in enumerate(data_int):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_page_break()
    
    # 六、视觉规范
    doc.add_heading('六、视觉规范', level=1)
    
    doc.add_heading('6.1 品牌色彩', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    colors = [
        ('用途', '色值', '说明'),
        ('品牌主色', '#8C1515', '北大红，用于标题、重点数字'),
        ('辅助色', '#1E3A5F', '深蓝，用于图表'),
        ('成功色', '#52C41A', '绿色，用于正向指标'),
        ('警告色', '#FAAD14', '橙色，用于预警'),
        ('背景色', '#F5F5F5', '浅灰，页面背景'),
        ('卡片色', '#FFFFFF', '白色，卡片背景'),
        ('文字主色', '#333333', '深灰，正文'),
    ]
    for i, row_data in enumerate(colors):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('6.2 字体规范', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    fonts = [
        ('用途', '字号', '字重'),
        ('KPI数字', '28px', 'Bold'),
        ('模块标题', '18px', 'SemiBold'),
        ('正文', '14px', 'Regular'),
        ('辅助文字', '12px', 'Regular'),
    ]
    for i, row_data in enumerate(fonts):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('6.3 圆角与间距', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    spacing = [
        ('元素', '规范'),
        ('卡片圆角', '12px'),
        ('按钮圆角', '8px'),
        ('卡片间距', '12px'),
    ]
    for i, (k, v) in enumerate(spacing):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_page_break()
    
    # 七、数据需求
    doc.add_heading('七、数据需求', level=1)
    
    doc.add_heading('7.1 数据来源', level=2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    sources = [
        ('数据', '来源表', '更新频率'),
        ('人员基本信息', '人员基本信息宽表', '每日'),
        ('编制信息', '编制信息表', '每月'),
        ('合同制信息', '合同制职工扩展和合同信息', '每日'),
        ('博士后信息', '博士后扩展信息', '每日'),
        ('退休预测', '在职退休时间计算', '每月'),
        ('人员进出', '合同制职工入职减离信息', '每日'),
        ('获奖信息', '独立获奖表（待提供）', '每月'),
    ]
    for i, row_data in enumerate(sources):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('7.2 数据权限', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    perms = [
        ('角色', '权限范围'),
        ('二级学院院长', '本学院全部数据 + 全校匿名对比'),
        ('学校管理员', '全校数据'),
    ]
    for i, (k, v) in enumerate(perms):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading('7.3 数据安全', level=2)
    security = [
        '手机号码等敏感信息需脱敏处理',
        '身份证号不展示',
        '全校对比不暴露其他学院具体数值',
    ]
    for item in security:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 八、非功能需求
    doc.add_heading('八、非功能需求', level=1)
    
    doc.add_heading('8.1 性能要求', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    perf = [
        ('指标', '要求'),
        ('首屏加载', '≤ 2秒'),
        ('图表渲染', '≤ 500ms'),
        ('数据刷新', '≤ 1秒'),
    ]
    for i, (k, v) in enumerate(perf):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading('8.2 兼容性要求', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    compat = [
        ('平台', '版本'),
        ('iOS', '≥ 12.0'),
        ('Android', '≥ 7.0'),
    ]
    for i, (k, v) in enumerate(compat):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_page_break()
    
    # 九、版本规划
    doc.add_heading('九、版本规划', level=1)
    
    versions = [
        ('V1.0（MVP）', [
            '人员概况KPI',
            '年龄、性别、学历、职称分布',
            '人员进出趋势',
            '人才名单',
            '全校对比排名',
        ]),
        ('V1.1', [
            '政治面貌分布',
            '学缘结构分析',
            '拟退预警',
            '获奖情况',
        ]),
        ('V1.2', [
            '交叉分析（职称-年龄、学历-年龄等）',
            '数据导出功能',
            '深色模式',
        ]),
    ]
    
    for ver, items in versions:
        doc.add_heading(ver, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 十、附录
    doc.add_heading('十、附录', level=1)
    
    doc.add_heading('10.1 数据字段映射', level=2)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    mapping = [
        ('展示指标', '数据字段', '说明'),
        ('性别', '性别', '加工数据'),
        ('出生日期', '出生日期', '原始数据'),
        ('政治面貌', '政治面貌', '原始数据'),
        ('最高学历', '文化程度（最后学历）中文', '加工数据'),
        ('职称', '聘任专业技术职务（聘）', '加工数据'),
        ('人员类别', '人员类别', '原始数据'),
        ('来校时间', '来校时间', '原始数据'),
        ('手机号码', '联系电话（手机）', '原始数据（需脱敏）'),
    ]
    for i, row_data in enumerate(mapping):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('10.2 视觉素材清单', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    assets = [
        ('素材', '用途'),
        ('北大校徽_红色.png', '顶部导航栏'),
        ('中英文校名_红色.png', '顶部导航栏'),
        ('标志与中英文校名组合规范_左右.png', '顶部导航栏'),
    ]
    for i, (k, v) in enumerate(assets):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    # 保存文档
    output_path = '/Users/liutao/Desktop/FR工作资料/项目/北京大学/学院看板移动端/PRD_学院看板移动端.docx'
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_doc()
